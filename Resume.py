import PyPDF2
import docx  # For DOCX
import os
import re
import logging
import json
from typing import List, Tuple, Optional, Dict
from collections import Counter
from transformers import pipeline
from datetime import datetime
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
import pandas as pd

# Download NLTK resources
try:
    nltk.download('punkt', quiet=True)
    nltk.download('stopwords', quiet=True)
except:
    pass

# Logging setup
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Global variables for models
summarizer = None
generator = None


# Initialize models with better error handling
def initialize_models():
    global summarizer, generator
    try:
        print("Loading models...")
        summarizer = pipeline("summarization", model="sshleifer/distilbart-cnn-12-6")
        generator = pipeline("text2text-generation", model="google/flan-t5-base")
        print("Transformers models loaded successfully")
    except Exception as e:
        logger.error(f"Transformers model loading failed: {e}")
        print(f"Warning: Model loading failed: {e}")
        summarizer = None
        generator = None


# Enhanced role-based skill sets with weights
ROLE_SKILLS = {
    "Data Analyst": {
        "core": ["SQL", "Python", "Excel", "Statistics", "Data Visualization"],
        "important": ["R", "Tableau", "Power BI", "Pandas", "NumPy"],
        "nice_to_have": ["SAS", "SPSS", "Jupyter", "Machine Learning", "ETL"]
    },
    "Project Manager": {
        "core": ["Project Planning", "Risk Management", "Team Leadership", "Communication"],
        "important": ["Agile", "Scrum", "Jira", "MS Project", "Stakeholder Management"],
        "nice_to_have": ["PMI", "PMP", "Kanban", "Budget Management", "Six Sigma"]
    },
    "Software Engineer": {
        "core": ["Programming", "Problem Solving", "Git", "Debugging"],
        "important": ["Java", "Python", "JavaScript", "System Design", "OOP"],
        "nice_to_have": ["Docker", "Kubernetes", "AWS", "Testing", "CI/CD"]
    },
    "Digital Marketing Specialist": {
        "core": ["SEO", "Content Marketing", "Social Media", "Analytics"],
        "important": ["Google Analytics", "PPC", "Email Marketing", "Facebook Ads"],
        "nice_to_have": ["A/B Testing", "Conversion Optimization", "Marketing Automation"]
    },
    "Web Developer": {
        "core": ["HTML", "CSS", "JavaScript", "Responsive Design"],
        "important": ["React", "Node.js", "API", "Database"],
        "nice_to_have": ["Vue.js", "Angular", "TypeScript", "GraphQL", "MongoDB"]
    },
    "DevOps Engineer": {
        "core": ["Docker", "CI/CD", "Cloud Platforms", "Monitoring"],
        "important": ["Kubernetes", "AWS", "Jenkins", "Terraform"],
        "nice_to_have": ["Ansible", "Prometheus", "ELK Stack", "Microservices"]
    },
}

# Industry keywords for better role detection
INDUSTRY_KEYWORDS = {
    "Data Science": ["machine learning", "deep learning", "neural networks", "AI", "data science"],
    "Finance": ["financial analysis", "investment", "portfolio", "risk assessment", "trading"],
    "Healthcare": ["medical", "clinical", "patient", "healthcare", "pharmaceutical"],
    "E-commerce": ["online retail", "e-commerce", "marketplace", "customer experience"],
    "Education": ["teaching", "curriculum", "student", "academic", "learning"]
}


logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_path: str) -> str:
    """Enhanced PDF text extraction using PyPDF2 with error handling and page numbering."""
    try:
        text = ""
        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}"
        return text.strip()
    except Exception as e:
        logger.error(f"Error reading PDF: {e}")
        return f"Error reading PDF: {e}"

def extract_text_from_docx(file_path: str) -> str:
    """Enhanced DOCX text extraction"""
    try:
        doc = docx.Document(file_path)
        full_text = []

        for para in doc.paragraphs:
            if para.text and para.text.strip():
                full_text.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        full_text.append(cell.text)

        return "\n".join(full_text).strip()
    except Exception as e:
        logger.error(f"Error reading DOCX: {e}")
        return f"Error reading DOCX: {e}"


def read_file(file) -> str:
    """Enhanced file reading with validation"""
    try:
        if not file:
            return "No file uploaded"

        if not hasattr(file, 'name') or not file.name:
            return "Invalid file object"

        if not os.path.exists(file.name):
            return "File not found"

        file_size = os.path.getsize(file.name)
        if file_size > 10 * 1024 * 1024:
            return "File too large (max 10MB)"

        if file_size == 0:
            return "File is empty"

        ext = os.path.splitext(file.name)[-1].lower()
        if ext == ".pdf":
            return extract_text_from_pdf(file.name)
        elif ext == ".docx":
            return extract_text_from_docx(file.name)
        else:
            return f"Unsupported file type: {ext}. Please upload PDF or DOCX files."
    except Exception as e:
        logger.error(f"Error reading file: {e}")
        return f"Error reading file: {str(e)}"


def extract_contact_info(text: str) -> Dict[str, str]:
    """Extract contact information using regex with better error handling"""
    contact_info = {}

    try:
        # Ensure text is a string
        if not isinstance(text, str):
            text = str(text) if text is not None else ""

        contact_info['name'] = extract_name_from_resume(text)

        # Email extraction
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        contact_info['email'] = emails[0] if emails else "Not found"

        # Phone extraction
        phone_patterns = [
            r'\+\d{1,3}[-.\s]?\d{3,4}[-.\s]?\d{3,4}[-.\s]?\d{4}',
            r'\(\d{3}\)\s?\d{3}[-.\s]?\d{4}',
            r'\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\b\d{10}\b',
            r'\d{3}\.\d{3}\.\d{4}',
            r'\d{3}\s\d{3}\s\d{4}',
            r'\+\d{1,3}\s?\d{3,4}\s?\d{3,4}\s?\d{4}',
            r'\d{4}[-.\s]?\d{3}[-.\s]?\d{3}',
        ]

        phone_found = None
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                phone_found = phones[0]
                break

        contact_info['phone'] = phone_found if phone_found else "Not found"

        # LinkedIn extraction
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        linkedin = re.findall(linkedin_pattern, text, re.IGNORECASE)
        contact_info['linkedin'] = linkedin[0] if linkedin else "Not found"

    except Exception as e:
        logger.error(f"Error extracting contact info: {e}")
        contact_info = {
            'name': "Error extracting name",
            'email': "Error extracting email",
            'phone': "Error extracting phone",
            'linkedin': "Error extracting LinkedIn"
        }

    return contact_info


def extract_name_from_resume(text: str) -> str:
    """Extract candidate name from resume text with better error handling"""
    try:
        # Ensure text is a string
        if not isinstance(text, str):
            text = str(text) if text is not None else ""

        if not text.strip():
            return "Name not found"

        lines = text.split('\n')

        for i, line in enumerate(lines[:5]):
            # Ensure line is a string
            if not isinstance(line, str):
                continue

            line = line.strip()
            if not line:
                continue

            skip_keywords = ['resume', 'curriculum vitae', 'cv', 'summary', 'objective',
                             'contact', 'profile', 'about', 'personal', 'details']

            # Safe keyword checking
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in skip_keywords):
                continue

            if line.isupper() and len(line.split()) >= 2:
                words = line.split()
                if 2 <= len(words) <= 4 and all(word.isalpha() for word in words if word):
                    return line.title()

            words = line.split()
            if len(words) >= 2 and all(word.isalpha() and len(word) > 0 and word[0].isupper() for word in words):
                false_positives = ['Bachelor Of', 'Master Of', 'Dear Sir', 'To Whom', 'Human Resources']
                if not any(fp.lower() in line_lower for fp in false_positives):
                    return line

    except Exception as e:
        logger.error(f"Error extracting name: {e}")

    return "Name not found"


def extract_experience_years(text: str) -> Optional[int]:
    """Extract years of experience from resume text"""
    try:
        if not isinstance(text, str):
            text = str(text) if text is not None else ""

        patterns = [
            r'(\d+)\+?\s*years?\s*of\s*experience',
            r'(\d+)\+?\s*years?\s*experience',
            r'experience\s*:?\s*(\d+)\+?\s*years?',
            r'(\d+)\+?\s*yrs?\s*exp'
        ]

        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                return int(matches[0])
    except Exception as e:
        logger.error(f"Error extracting experience years: {e}")

    return None


def extract_education(text: str) -> List[str]:
    """Extract education information"""
    try:
        if not isinstance(text, str):
            text = str(text) if text is not None else ""

        education_keywords = [
            'bachelor', 'master', 'phd', 'doctorate', 'diploma', 'certificate',
            'b.tech', 'b.sc', 'm.tech', 'm.sc', 'mba', 'bca', 'mca'
        ]

        education_info = []
        lines = text.split('\n')

        for line in lines:
            if not isinstance(line, str):
                continue

            line_lower = line.lower()
            if any(keyword in line_lower for keyword in education_keywords):
                education_info.append(line.strip())

        return education_info[:3]
    except Exception as e:
        logger.error(f"Error extracting education: {e}")
        return []


def detect_job_role_from_text(text: str) -> Tuple[Optional[str], float]:
    """Enhanced job role detection with confidence score"""
    try:
        if not isinstance(text, str):
            text = str(text) if text is not None else ""

        text_lower = text.lower()
        role_scores = {}

        for role in ROLE_SKILLS.keys():
            role_lower = role.lower()
            if role_lower in text_lower:
                role_scores[role] = text_lower.count(role_lower) * 10

        for role, skills_dict in ROLE_SKILLS.items():
            score = 0
            all_skills = skills_dict['core'] + skills_dict['important'] + skills_dict['nice_to_have']

            for skill in all_skills:
                if skill.lower() in text_lower:
                    if skill in skills_dict['core']:
                        score += 3
                    elif skill in skills_dict['important']:
                        score += 2
                    else:
                        score += 1

            role_scores[role] = role_scores.get(role, 0) + score

        if role_scores:
            best_role = max(role_scores, key=role_scores.get)
            confidence = min(100, role_scores[best_role] * 2)
            return best_role, confidence

    except Exception as e:
        logger.error(f"Error detecting job role: {e}")

    return None, 0


def advanced_skill_scoring(resume_text: str, required_skills: Dict[str, List[str]]) -> Tuple[
    float, str, List[str], Dict[str, int]]:
    """Advanced skill scoring with weighted categories"""
    try:
        if not isinstance(resume_text, str):
            resume_text = str(resume_text) if resume_text is not None else ""

        found_skills = []
        skill_counts = {}
        resume_lower = resume_text.lower()

        core_score = 0
        important_score = 0
        nice_score = 0

        total_core_skills = len(required_skills.get('core', []))
        total_important_skills = len(required_skills.get('important', []))
        total_nice_skills = len(required_skills.get('nice_to_have', []))

        for skill in required_skills.get('core', []):
            skill_lower = skill.lower()
            if skill_lower in resume_lower:
                found_skills.append(skill)
                count = resume_lower.count(skill_lower)
                skill_counts[skill] = count
                core_score += 3

        for skill in required_skills.get('important', []):
            skill_lower = skill.lower()
            if skill_lower in resume_lower:
                found_skills.append(skill)
                count = resume_lower.count(skill_lower)
                skill_counts[skill] = count
                important_score += 2

        for skill in required_skills.get('nice_to_have', []):
            skill_lower = skill.lower()
            if skill_lower in resume_lower:
                found_skills.append(skill)
                count = resume_lower.count(skill_lower)
                skill_counts[skill] = count
                nice_score += 1

        total_possible_score = (total_core_skills * 3) + (total_important_skills * 2) + (total_nice_skills * 1)
        actual_score = core_score + important_score + nice_score

        percentage_score = (actual_score / total_possible_score) * 100 if total_possible_score > 0 else 0

        if percentage_score >= 80:
            level = "Excellent Match"
        elif percentage_score >= 60:
            level = "Strong Match"
        elif percentage_score >= 40:
            level = "Good Match"
        elif percentage_score >= 20:
            level = "Partial Match"
        else:
            level = "Needs Improvement"

        return min(100, percentage_score), level, found_skills, skill_counts

    except Exception as e:
        logger.error(f"Error in skill scoring: {e}")
        return 0, "Error in Analysis", [], {}


def generate_detailed_feedback(score: float, found: List[str], required: Dict[str, List[str]],
                               contact_info: Dict[str, str], experience_years: Optional[int]) -> str:
    """Generate comprehensive feedback"""
    try:
        feedback_parts = []

        if score >= 80:
            feedback_parts.append("🎉 Excellent match for this role!")
        elif score >= 60:
            feedback_parts.append("✅ Good alignment with role requirements.")
        else:
            feedback_parts.append("⚠️ Resume needs improvement to match role requirements.")

        all_required = required.get('core', []) + required.get('important', []) + required.get('nice_to_have', [])
        missing_core = [s for s in required.get('core', []) if s not in found]
        missing_important = [s for s in required.get('important', []) if s not in found]

        if missing_core:
            feedback_parts.append(f"🔴 Critical skills missing: {', '.join(missing_core[:3])}")

        if missing_important:
            feedback_parts.append(f"🟡 Important skills to add: {', '.join(missing_important[:3])}")

        contact_issues = []
        if contact_info.get('email', '') == "Not found":
            contact_issues.append("email")
        if contact_info.get('phone', '') == "Not found":
            contact_issues.append("phone")
        if contact_info.get('linkedin', '') == "Not found":
            contact_issues.append("LinkedIn profile")

        if contact_issues:
            feedback_parts.append(f"📧 Add missing contact info: {', '.join(contact_issues)}")

        if experience_years:
            feedback_parts.append(f"💼 Experience: {experience_years} years detected")
        else:
            feedback_parts.append("💼 Consider clearly stating years of experience")

        return "\n".join(feedback_parts)

    except Exception as e:
        logger.error(f"Error generating feedback: {e}")
        return f"Error generating feedback: {str(e)}"


def generate_ats_score(resume_text: str) -> Tuple[int, List[str]]:
    """Generate comprehensive ATS compatibility score"""
    try:
        if not isinstance(resume_text, str):
            resume_text = str(resume_text) if resume_text is not None else ""

        ats_issues = []
        score = 100

        email_found = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', resume_text)
        if not email_found:
            ats_issues.append("Email address missing")
            score -= 15

        phone_patterns = [
            r'\+\d{1,3}[-.\s]?\d{3,4}[-.\s]?\d{3,4}[-.\s]?\d{4}',
            r'\(\d{3}\)\s?\d{3}[-.\s]?\d{4}',
            r'\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\b\d{10}\b',
            r'\d{3}\.\d{3}\.\d{4}',
            r'\d{3}\s\d{3}\s\d{4}',
        ]

        phone_found = any(re.search(pattern, resume_text) for pattern in phone_patterns)
        if not phone_found:
            ats_issues.append("Phone number missing or poorly formatted")
            score -= 10

        word_count = len(resume_text.split())
        if word_count < 200:
            ats_issues.append("Resume too short (less than 200 words)")
            score -= 15
        elif word_count > 1000:
            ats_issues.append("Resume too long (over 1000 words)")
            score -= 5

        special_chars = len(re.findall(r'[^\w\s\-\.\,\(\)\@]', resume_text))
        if special_chars > len(resume_text) * 0.05:
            ats_issues.append("Too many special characters")
            score -= 10

        sections_to_check = [
            (r'\bexperience\b', "Experience section"),
            (r'\beducation\b', "Education section"),
            (r'\bskills?\b', "Skills section"),
        ]

        for pattern, section_name in sections_to_check:
            if not re.search(pattern, resume_text, re.IGNORECASE):
                ats_issues.append(f"{section_name} not clearly marked")
                score -= 8

        date_patterns = [
            r'\b\d{4}\s*[-–]\s*\d{4}\b',
            r'\b\d{1,2}/\d{4}\s*[-–]\s*\d{1,2}/\d{4}\b',
            r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+\d{4}\b',
            r'\b\d{4}\s*[-–]\s*present\b',
        ]
        date_found = any(re.search(pattern, resume_text, re.IGNORECASE) for pattern in date_patterns)
        if not date_found:
            ats_issues.append("Employment dates not clearly formatted")
            score -= 8

        return max(0, score), ats_issues

    except Exception as e:
        logger.error(f"Error generating ATS score: {e}")
        return 0, [f"Error analyzing ATS compatibility: {str(e)}"]


def simple_summarize(text: str) -> str:
    """Simple text summarization fallback"""
    try:
        if not isinstance(text, str):
            text = str(text) if text is not None else ""

        sentences = text.split('.')
        if len(sentences) > 3:
            return '. '.join(sentences[:2] + [sentences[-1]]).strip()
        return text[:200] + "..." if len(text) > 200 else text
    except Exception as e:
        logger.error(f"Error in simple summarize: {e}")
        return "Error generating summary"


def validate_resume_content(text: str) -> bool:
    """Validate if the content appears to be a resume by checking for key keywords"""
    resume_keywords = ["experience", "education", "skills", "contact", "work history", "qualifications"]

    # Check if at least one resume keyword is present
    for keyword in resume_keywords:
        if keyword.lower() in text.lower():
            return True

    return False


def analyze_resume(resume_text: str, jd_text: str = "", target_role: str = "Auto-detect") -> Dict[str, any]:
    """Enhanced resume analysis function returning JSON-serializable dict"""
    try:
        if not resume_text or len(resume_text) < 50:
            return {"error": "Invalid or unreadable resume."}

        # Validate if content appears to be a resume
        if not validate_resume_content(resume_text):
            return {"error": "The uploaded file does not appear to be a resume. Please upload a valid resume document containing sections like experience, education, skills, or qualifications."}

        contact_info = extract_contact_info(resume_text)
        experience_years = extract_experience_years(resume_text)
        education = extract_education(resume_text)

        if jd_text:
            detected_role, confidence = detect_job_role_from_text(jd_text)
        else:
            detected_role, confidence = detect_job_role_from_text(resume_text)

        if target_role and target_role != "Auto-detect":
            detected_role = target_role
            # Calculate confidence based on skill match score
            required_skills_for_confidence = ROLE_SKILLS.get(detected_role, {})
            if required_skills_for_confidence:
                confidence_score, _, _, _ = advanced_skill_scoring(resume_text, required_skills_for_confidence)
                confidence = confidence_score
            else:
                confidence = 0

        if detected_role and detected_role in ROLE_SKILLS:
            required_skills = ROLE_SKILLS[detected_role]
        else:
            required_skills = {
                "core": ["Communication", "Teamwork", "Problem Solving"],
                "important": ["Leadership", "Time Management", "Adaptability"],
                "nice_to_have": ["Innovation", "Customer Service", "Technical Skills"]
            }

        score, level, found_skills, skill_counts = advanced_skill_scoring(resume_text, required_skills)
        feedback = generate_detailed_feedback(score, found_skills, required_skills, contact_info, experience_years)
        ats_score, ats_issues = generate_ats_score(resume_text)

        try:
            if summarizer:
                summary = summarizer(resume_text[:1000], max_length=150, min_length=30, do_sample=False)[0][
                    "summary_text"]
            else:
                summary = simple_summarize(resume_text)
        except Exception as e:
            print(f"Summarization error: {e}")
            summary = simple_summarize(resume_text)

        skill_breakdown = f"\n\n📊 Skill Analysis:\n"
        skill_breakdown += f"• Core skills found: {len([s for s in found_skills if s in required_skills.get('core', [])])}/{len(required_skills.get('core', []))}\n"
        skill_breakdown += f"• Important skills found: {len([s for s in found_skills if s in required_skills.get('important', [])])}/{len(required_skills.get('important', []))}\n"
        skill_breakdown += f"• Nice-to-have skills found: {len([s for s in found_skills if s in required_skills.get('nice_to_have', [])])}/{len(required_skills.get('nice_to_have', []))}"

        enhanced_feedback = feedback + skill_breakdown

        return {
            "summary": summary,
            "score": f"{score:.1f}% ({level})",
            "role": f"{detected_role or 'Not Detected'} (Confidence: {confidence:.1f}%)",
            "skills": ", ".join(found_skills) if found_skills else "❌ No matching skills found",
            "feedback": enhanced_feedback,
            "contact": {
                "name": contact_info.get('name', 'N/A'),
                "email": contact_info.get('email', 'N/A'),
                "phone": contact_info.get('phone', 'N/A'),
                "linkedin": contact_info.get('linkedin', 'N/A')
            },
            "education": [f"🎓 {edu}" for edu in education] if education else ["❌ No education information found"],
            "ats": {
                "score": ats_score,
                "issues": ats_issues
            }
        }

    except Exception as e:
        error_msg = f"Error analyzing resume: {str(e)}"
        logger.error(error_msg)
        print(error_msg)
        return {"error": error_msg}


# Initialize models when module is imported (optional)
# Commented out to avoid memory issues on deployment
# try:
#     initialize_models()
# except Exception as e:
#     print(f"Warning: Could not initialize models: {e}")
#     print("Continuing without AI models - basic functionality will work")

